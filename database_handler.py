## modBot Project
## database_handler.py

import os
from typing import List, Dict, Any, Optional
import logging
from openai import OpenAI
from config import get_config, Config
from astrapy import DataAPIClient
from astrapy.constants import VectorMetric, SortDocuments
from openai import OpenAI
import datetime
from dotenv import load_dotenv
from langchain_community.document_loaders import WebBaseLoader
from langchain_text_splitters import RecursiveCharacterTextSplitter
import requests
import rtfparse
from bs4 import BeautifulSoup
import time
import asyncio
from concurrent.futures import ThreadPoolExecutor
import json
import warnings 
import docx
import pandas as pd
import PyPDF2
import rtfparse


#########################################################################################

### LOAD VARIABLES
load_dotenv()
# Set .env variables
api_key = os.getenv("OPENAI_API_KEY")
chat_model = os.getenv("CHAT_MODEL")
vector_model = os.getenv("EMBEDDING_MODEL")
token = os.getenv("ASTRA_DB_TOKEN")
api_endpoint = os.getenv("ASTRA_DB_API_ENDPOINT")
USER_AGENT = os.getenv("USER_AGENT")
keyspace = os.getenv("ASTRA_DB_KEYSPACE")
collection_name = os.getenv("ASTRA_DB_COLLECTION")
dimension = int(os.getenv("VECTOR_DIMENSION"))
user_name = os.getenv("USERNAME")
#########################################################################################

# Suppress deprecation warnings
warnings.filterwarnings("ignore", category=DeprecationWarning)

# Set up logging
#logging.basicConfig(level=logging.CRITICAL)
#logging.basicConfig(level=logging.DEBUG)
#logging.basicConfig(level=logging.INFO)
logging.disable(level=logging.DEBUG)
logger = logging.getLogger(__name__)

#########################################################################################

class DatabaseHandler:
    def __init__(self, token, api_endpoint, keyspace, collection_name, recency_collection_name, dimension, chat_service):
        """Initialize database connection and embeddings."""
        self.config = get_config()
        self.token = token
        self.api_endpoint = api_endpoint
        self.keyspace = keyspace
        self.collection_name = collection_name
        self.recency_collection_name = recency_collection_name  # Store the recency collection name
        self.dimension = dimension
        self.chat_service = chat_service  # Add this line to store the ChatService instance
        
        
        # Initialize vector client
        self.vector_client = OpenAI()
        
        # Initialize database client
        self.db_client = DataAPIClient(self.token)
        
        # Initialize database
        self.db = self.db_client.get_database_by_api_endpoint(self.api_endpoint, keyspace=self.keyspace)
        
        self.collection = self.create_collection(self.collection_name, self.dimension)
        self.recency_collection = self.create_collection(self.recency_collection_name, self.dimension)  # Create the recency collection
        
        logger.info("Connected to Astra DB")

############################################****NEW TEST CODE****###############################################################

    async def get_sorted_recent_conversation(self):
        logger.debug("Retrieving documents from the recency collection.")
        
        cursor = self.recency_collection.find({}, limit=15)  # Retrieve documents
        recent_docs = await asyncio.to_thread(list, cursor)  # Convert cursor to list using a separate thread

        logger.debug(f"Recent Conversation documents retrieved: {recent_docs}")
        #print(f"Recent Conversation documents retrieved: {recent_docs}")

        # Strip the embedding before sorting
        for doc in recent_docs:
            if '$vector' in doc:
                logger.debug(f"Removing embedding from document: {doc}")
                del doc['$vector']
        #print(f"Recent conversation documents stripped of embeddings: {recent_docs}")

        # Sort the documents in ascending order based on timestamp
        recent_docs.sort(key=lambda x: x['timestamp'])

        logger.debug(f"Retrieved and sorted documents without embeddings: {recent_docs}")
        return recent_docs

############################################****END OF NEW TEST CODE****########################################################

    def create_collection(self, collection_name, dimension):
        """Create a new collection in the database."""
        collection = self.db.create_collection(
            collection_name,
            dimension=dimension,
            metric=VectorMetric.COSINE,
            check_exists=False
        )
        logger.info(f"* Collection: {collection.full_name}\n")
        return collection
        
    def get_embedding(self, text):
        logger.debug(f"\nArrived in embedding function for: {text}")
        """Generate embedding for a given text."""
        response = self.vector_client.embeddings.create(
            input=text,
            ## fix this so model=model from dotenv
            model="text-embedding-3-small"
        )

        embedding = response.data[0].embedding

        if len(embedding) > self.dimension:
            embedding = embedding[:self.dimension]
        elif len(embedding) < self.dimension:
            embedding += [0] * (self.dimension - len(embedding))
        #print(f"\n Requested embedding########################: {embedding}")
        return embedding

    async def delete_all_documents_of_type_document(self, del_type):
        # Find all documents with the specified type
        documents_to_delete = self.collection.find({"type": del_type})
        
        deleted_count = 0
        for document in documents_to_delete:
            self.collection.delete_one({"_id": document["_id"]})
            deleted_count += 1
            
        return deleted_count

    async def insert_message(self, document):
        logger.debug(f"\n Arrived in insert_message with document type: {type(document)}")
        document = document
        # Check if the document type is "external_document"
        logger.debug("\n Arrived in insert_message")  
        #print(type(document))
        #print(f" \nCombined document arrived in insert_message: {document}")
        if document.get("type") == "external_document":
            # Calculate expiration timestamp (1 week later)
            expiration_timestamp = datetime.datetime.now() + datetime.timedelta(weeks=5)
            document["expiration"] = expiration_timestamp.isoformat()  # Add expiration field
            
        # Add identifier for past conversations
        if "is_read_from_external_document" in document:
            document["external_document"] = True  # Indicates it is from past conversations

        # Add identifier for past conversations
        if "is_read_from_url" in document:
            document["url"] = True  # Indicates it is from past conversations

        # Add identifier for past conversations
        if "is_from_past_conversations" in document:
            document["from_past_conversations"] = True  # Indicates it is from past conversations

        self.collection.insert_one(document)
        logger.debug(f"\nDocument inserted into database {document}")

        # Remove the embedding from the document
        if '$vector' in document:
            del document['$vector']
            logger.debug("\nEmbedding removed from combined document.")

        # Conditional Return
        if document.get("type") in ["external_document", "url"]:
            return

        # Change identifier for the recency collection
        if "from_past_conversations" in document:
            document["is_from_current_conversation_session"] = True  # Change identifier for recency


        # Insert the same document into the recency collection
        self.recency_collection.insert_one(document)  # Insert the same document into the recency collection
        logger.debug(f"\nDocument inserted into recency collection: {document}")

        # Maintain the recency collection to limit its size of 5 pairs
        await self.maintain_recency_collection()  # Call to maintain the size of the recency collection

    async def store_url(self, document_text):
        logger.debug("Arrived in store_url")
        # Chunk the document
        logger.debug("\nBeginning document chunking in url function")
        splitter = RecursiveCharacterTextSplitter(
            chunk_size=1460,
            chunk_overlap=30
        )
        chunks = splitter.split_text(document_text)
        logger.debug("\nDocument chunks created")
        
        logger.debug("\nAdd type 'url' to each chunk before storing in db")
        for chunk in chunks:
            embedding = self.get_embedding(chunk)
            document = {"type": "url", "text": chunk, "$vector": embedding, "timestamp": datetime.datetime.now().isoformat()}
            logger.debug("\nSending url chunks for storage in db")
            await self.insert_message(document)
            logger.debug("\nReturning from storing url document chunk")
        print("Website scraping chunks added to DB.")

    async def add_external_documents(self, file_path):
        print("\n Arrived in add_external_documents")
        if not os.path.isfile(file_path):
            print("File not found. Please check the path and try again.")
            return

        file_extension = os.path.splitext(file_path)[1].lower()
        document_text = ""
        print("File extension is:  ", file_extension)

        # Read content based on file type
        try:
            if file_extension == '.txt':
                with open(file_path, 'r', encoding='utf-8') as file:
                    document_text = file.read()

            elif file_extension == '.docx':
                doc = docx.Document(file_path)
                document_text = '\n'.join([para.text for para in doc.paragraphs])

            elif file_extension == '.pdf':
                with open(file_path, 'rb') as file:
                    reader = PyPDF2.PdfReader(file)
                    document_text = '\n'.join([page.extract_text() for page in reader.pages if page.extract_text()])
                    print("PDF doc text is: ", document_text)

            elif file_extension == '.rtf':
                with open(file_path, 'r', encoding='utf-8') as file:
                    document_text = rtfparse.parse(file.read())

            elif file_extension == '.xlsx':
                df = pd.read_excel(file_path)
                document_text = df.to_string(index=False)

            elif file_extension == '.csv':
                df = pd.read_csv(file_path)
                document_text = df.to_string(index=False)

            else:
                print("Unsupported file type. Please provide a .txt, .docx, .pdf, .rtf, or .xlsx file.")
                return

        except Exception as e:
            print(f"An error occurred while reading the file: {e}")
            return
        print("Document text read is: ", document_text)

        # Chunk the document
        print("\nBeginning document chunking in external documents function")
        splitter = RecursiveCharacterTextSplitter(
            chunk_size=1460,
            chunk_overlap=30
        )
        chunks = splitter.split_text(document_text)
        print("document text: ", document_text)
        print("\nDocument chunks created")

        print("\nAdd type 'external_document' to each chunk before storing in db")
        for chunk in chunks:
            print("this chunk is : ", chunk)
            embedding = self.get_embedding(chunk)
            document = {"type": "external_document", "text": chunk, "$vector": embedding, "timestamp": datetime.datetime.now().isoformat()}
            print("\nSending external document chunks for storage in db")
            await self.insert_message(document)
            print("\nReturning from storing external document chunk")
        print("Document chunks added to DB.")



    async def semantic_search(self, query: str, k = 3) -> List[Dict]:
        logger.debug(f"Starting semantic search. Query: '{query}'")

        # Create a result list to accumulate results
        combined_results = []
        unique_texts = set()  # Set to track unique document texts
        logger.debug(f"\n combined_results and unique_texts initialized as blank dictionary and list, respectively. \n combined_results: {combined_results} \n unique_texts: {unique_texts}")

        # Retrieve and log sorted recent documents
        logger.debug("Starting recent_convo search")
        last_five_convo = await self.get_sorted_recent_conversation()
        logger.debug(f"Sorted recent conversation documents: {last_five_convo}")
        #print(f"Sorted recent conversation documents: {last_five_convo}")

        
        # Add last_five_convo to combined_results with an identifier
        for convo in last_five_convo:
            convo["is_from_current_conversation_session"] = True  # Add identifier for recent conversation
            combined_results.append(convo)
        logger.debug(f"Added 'is_from_current_conversation_session' to each documment: {combined_results}")
        #print(f"Added 'is_from_current_conversation_session' to each documment: {combined_results}")

        # Generate embedding for the query
        #print(f"Query: {query}")
        human_embedding = self.get_embedding(query)
        #print(f"Human embedding: {human_embedding}")
        #human_document = {"type": ["human"], "text": query, "$vector": human_embedding, "timestamp": datetime.datetime.now().isoformat()}
        #await self.insert_message(human_document)

        # Similarity search in conversation messages
        logger.debug("\nPerforming search of conversation history.")
        k = 15
        conversation_results = list(self.collection.find(
        {"type": {"$in": ["conversation"]}},  # Filter by type "conversation"
        sort={"$vector": human_embedding},   # Ensure correct vector sorting
        limit=k,
        ))
        logger.debug(f"\n Direct results of conversation search {conversation_results}")

        # Deduplicate conversation results
        for result in conversation_results:
            if result['text'] not in unique_texts:  # Check if the text is unique
                unique_texts.add(result['text'])  # Add text to the set
                combined_results.append(result)  # Add to the result list
 
        # Add conversation_results to combined_results with an identifier for historical conversations
        for convo in conversation_results:
            convo["is_from_past_conversations"] = True  # Add identifier for historical conversation
            combined_results.append(convo)
        logger.debug(f"\nHistorical conversation results added to context variable: {combined_results}")
        
        # Similarity search in uploaded documents
        logger.debug("Performing similarity search in uploaded documents.")
        k = 25
        doc_results = list(self.collection.find(
            {"$or": [ {"type": "external_document"}, {"type": "url"} ]},  # Filter by types "external_document" or "url"
            sort={"$vector": human_embedding},   # Ensure correct vector sorting
            limit=k,
        ))
        logger.debug(f"\n Direct results of uploaded docs search: {doc_results}")

        # Deduplicate document results and add identifier
        for result in doc_results:
            if result['text'] not in unique_texts:  # Check if the text is unique
                unique_texts.add(result['text'])  # Add text to the set
                result["is_read_from_external_document"] = True  # Add identifier for external document
                combined_results.append(result)  # Add to the result list

        logger.debug(f"\n combined_results updated with deduplicated documents: {combined_results}")
        
        ## log successfully appending non-duplicate results to combined_results
        logger.debug(f"\nCombined results of conversation search and uploaded docs search: {combined_results}. Type: {type(combined_results)}")
        return combined_results  # Return combined results from both searches



    async def maintain_recency_collection(self):
        logger.debug("\nArrived in maintain_recency_collection")
        
        # Start a loop to maintain the recency collection
        while True:
            # Count documents in the recency collection with an upper bound limit
            count = self.recency_collection.count_documents({}, upper_bound=30)  # Use an appropriate number for upper bound
            logger.debug(f"\nCurrent count of documents in recency collection: {count}")

            # Check if the count exceeds the limit of 10
            if count <= 15:
                logger.debug("No need to delete any documents, count is within limit.")
                break  # Exit the loop if the count is acceptable
            
            # Find the oldest document based on timestamp and delete it
            oldest_document = self.recency_collection.find_one(
                {},  # No filter
                sort={"timestamp": SortDocuments.ASCENDING}  # Sort by timestamp ascending
            )
            
            if oldest_document:
                self.recency_collection.delete_one({"_id": oldest_document["_id"]})  # Delete the oldest document
                logger.debug(f"\nDeleted oldest document from recency collection: {oldest_document}")
            else:
                logger.debug("No oldest document found to delete. Exiting maintain_recency_collection.")
                break  # If there are no documents, exit the loop



    async def keyword_search(self, query: str, k: int = 5) -> List[Dict]:
        # Sample query for testing
        query = query
        logger.debug(f"Performing keyword search for query: '{query}'")

        # Prepare the prompt for keyword extraction returning a single string
        user_input = (
            "You are an AI model designed to extract relevant keywords from text. "
            "Your task is to identify and return the most important keywords found in the given input text. "
            "Please ensure the keywords are relevant and accurately represent the main topics of the text.\n"
            f"Input: {query}\n"
            "Output: Please return a single string containing the keywords separated by spaces, e.g., 'keyword1 keyword2 keyword3'.\n"
        )

        logger.debug(f"\nPrompt for keyword search query: {user_input} Type: {type(user_input)}")

        logger.debug("\n Calling the chat method to extract keywords.")

        ## Perform keywords search of database
        keywords_response = await self.chat_service.chat(user_input)
        logger.debug(f"\nKeywords extracted from query: {keywords_response} \n Type: {type(keywords_response)}")
        logger.error(f"Error extracting keywords from chat response: {e}")
        return []
     
        ## Log and print statements
        logger.debug(f"Query for keyword search logged: '{query}', resulting in the following keywords: {keywords_response}")

        # Perform semantic search for keywords
        logger.debug("Starting keyword search in the database.")
        query = keywords_response
        logger.debug(f"\n For keyword semantic search: Query: {query}")
        results = await self.semantic_search(query, search_type, k=3)
        keyword_search_results = results

        logger.debug("Keyword search completed.")

        # Debug: Checking how many results were found
        if keyword_search_results is not None:
            logger.debug(f"Keyword search completed. Results: {keyword_search_results}")
        else:
            logger.warning("No results found for the query.")
        
        # Log the results
        if keyword_search_results and len(keyword_search_results) > 0:
            logger.debug(f"Found {len(keyword_search_results)} results.")
            for idx, result in enumerate(keyword_search_results):
                logger.debug(f"Result {idx + 1}: {keyword_search_results}")
                # If the result contains embedding info, log it
                if '$vector' in keyword_search_results:
                    logger.debug(f"Embedding for result {idx + 1}: {keyword_search_results['$vector']}")
        else:
            logger.warning("No results found for the query.")



    async def clear_documents(self):
        """Clear all documents from vector store"""
        # Adjusted to use the new database connection
        # Add the logic for clearing documents here
        pass

class Document:
    def __init__(self, page_content, metadata):
        self.page_content = page_content
        self.metadata = metadata

    def __repr__(self):
        return f"Document(text={self.page_content}, metadata={self.metadata})"

    def __str__(self):
        return ' '.join(self.page_content.split())  # Replace multiple spaces/newlines with a single space